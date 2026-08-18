"""Upload de documento livre com auto-preenchimento das abas via LLM opcional.

Suporta PDF, DOCX e TXT. Duas branches:

- Com chave Groq (BYOK — cada aluno cadastra a sua em console.groq.com):
  usa Llama 3.3 70B via API compatível OpenAI para extrair contexto rico,
  dor, dados, riscos, valor e casos de uso da Aula 2.
- Sem chave: fallback heurístico determinístico (regex + keywords) que
  extrai o que consegue e devolve o texto bruto para o aluno colar
  manualmente nos campos.

Prompt do LLM força saída em JSON estruturado para não quebrar o pipeline.
"""

from __future__ import annotations

import io
import json
import os
import re
from typing import Any

from src.priorizacao import CRITERIO_IDS, CasoDeUso

MAX_CHARS_LLM = 12000
MODEL_DEFAULT = "llama-3.3-70b-versatile"
GROQ_BASE_URL = "https://api.groq.com/openai/v1"
STATE_KEY_GROQ = "groq_api_key"


class ErroExtracao(Exception):
    pass


def extrair_texto(arquivo_bytes: bytes, nome_arquivo: str) -> str:
    """Extrai texto puro de PDF, DOCX ou TXT."""
    nome = nome_arquivo.lower()
    if nome.endswith(".pdf"):
        return _extrair_pdf(arquivo_bytes)
    if nome.endswith(".docx"):
        return _extrair_docx(arquivo_bytes)
    if nome.endswith(".txt") or nome.endswith(".md"):
        return arquivo_bytes.decode("utf-8", errors="replace")
    raise ErroExtracao(f"Formato não suportado: {nome_arquivo}. Use PDF, DOCX, TXT ou MD.")


def _extrair_pdf(dados: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(dados))
    partes = []
    for page in reader.pages:
        try:
            partes.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n\n".join(partes).strip()


def _extrair_docx(dados: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(dados))
    paragrafos = [p.text for p in doc.paragraphs if p.text.strip()]
    # tabelas viram texto plano
    for tabela in doc.tables:
        for linha in tabela.rows:
            cells = [c.text.strip() for c in linha.cells if c.text.strip()]
            if cells:
                paragrafos.append(" | ".join(cells))
    return "\n".join(paragrafos).strip()


def api_key_disponivel() -> bool:
    """Verifica se há chave Groq (env var GROQ_API_KEY, st.secrets ou
    session_state do aluno)."""
    return _resolver_api_key() is not None


def _resolver_api_key() -> str | None:
    """Resolve a chave Groq em ordem: session_state (BYOK do aluno) → env
    var → st.secrets. Retorna None se nenhuma fonte tem chave."""
    try:
        import streamlit as st

        key = (st.session_state.get(STATE_KEY_GROQ) or "").strip()
        if key:
            return key
    except Exception:
        pass
    key = os.getenv("GROQ_API_KEY")
    if key:
        return key
    try:
        import streamlit as st

        if hasattr(st, "secrets") and "GROQ_API_KEY" in st.secrets:
            return st.secrets["GROQ_API_KEY"]
    except Exception:
        pass
    return None


PROMPT_SISTEMA = """Você é um analista sênior de PPPM (Portfólio, Programa e Projeto) e IA aplicada, treinado no método do Prof. Dr. José Bezerra.

Leia o documento fornecido e extraia informações estruturadas para preencher o app consultor-ia-pppm.

Regras rígidas:
- Não invente empresas, nomes ou números que não estão no documento.
- Se um campo não puder ser inferido com segurança, retorne string vazia.
- Casos de uso propostos DEVEM se apoiar em dor/dado citados no próprio documento.
- Ranking, notas e priorização SÃO responsabilidade do aluno depois — não invente notas.
- Português (pt-BR) em todas as strings.

Retorne APENAS JSON válido no schema abaixo, sem nenhum texto antes ou depois.
"""

SCHEMA_JSON = """{
  "contexto": {
    "empresa": "string — nome da organização (vazio se não citada)",
    "porte": "MEI | Pequena | Média | Grande | N/A",
    "cargo": "string — cargo do responsável se citado (vazio caso contrário)",
    "n_projetos": "int — quantidade de projetos ativos citados, ou 0",
    "pmo_ativo": "bool — true se existe PMO/escritório de projetos citado"
  },
  "mapa": {
    "contexto": "string — resumo do contexto de PPPM em 2-4 frases",
    "dor": "string — dor principal declarada em 1-2 frases objetivas",
    "dados": "string — sistemas/fontes de dados disponíveis citados",
    "riscos": "string — riscos declarados no documento",
    "valor": "string — resultado esperado ou métrica de valor mencionada"
  },
  "casos_uso": [
    {
      "id": "string curto tipo caso-1",
      "nome": "string — título objetivo do caso",
      "contexto": "projeto|programa|portfólio|PMO",
      "dor": "string — problema real que este caso resolve",
      "dados": "string — dados/fontes necessários",
      "decisao": "string — qual escolha será melhorada",
      "dono": "string — cargo do dono humano da decisão, ou vazio",
      "metrica_valor": "string — como o benefício será medido"
    }
  ]
}"""


def sugerir_via_llm(texto: str, api_key: str | None = None, model: str = MODEL_DEFAULT) -> dict:
    """Chama o Groq (via SDK OpenAI-compatível) e devolve sugestões estruturadas.

    Levanta ErroExtracao se a chamada falhar ou o JSON de retorno for inválido.
    """
    from openai import OpenAI

    key = api_key or _resolver_api_key()
    if not key:
        raise ErroExtracao("Chave Groq não configurada.")
    client = OpenAI(api_key=key, base_url=GROQ_BASE_URL)
    conteudo = texto[:MAX_CHARS_LLM]
    prompt_user = (
        f"Schema esperado (JSON):\n{SCHEMA_JSON}\n\n"
        f"Documento (limitado a {MAX_CHARS_LLM} caracteres):\n---\n{conteudo}\n---"
    )
    try:
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": PROMPT_SISTEMA},
                {"role": "user", "content": prompt_user},
            ],
            response_format={"type": "json_object"},
            temperature=0.2,
        )
        raw = resp.choices[0].message.content or "{}"
        return json.loads(raw)
    except json.JSONDecodeError as e:
        raise ErroExtracao(f"LLM devolveu JSON inválido: {e}") from e
    except Exception as e:
        raise ErroExtracao(f"Falha na chamada LLM (Groq): {e}") from e


def heuristica_fallback(texto: str) -> dict:
    """Extração determinística sem LLM — heurísticas simples.

    Retorna a mesma estrutura de `sugerir_via_llm` mas com campos vazios se
    não conseguir inferir. Usado quando a chave OpenAI não está configurada.
    """
    t = texto or ""
    contexto = {"empresa": _achar_empresa(t), "porte": "", "cargo": "", "n_projetos": 0, "pmo_ativo": _tem_pmo(t)}
    mapa = {
        "contexto": _resumo_primeiras_frases(t, n=3),
        "dor": _primeira_ocorrencia_apos(t, ["problema", "dor", "gargalo", "dificuldade"]),
        "dados": _primeira_ocorrencia_apos(t, ["dados", "sistema", "base", "planilha", "jira", "erp", "crm"]),
        "riscos": _primeira_ocorrencia_apos(t, ["risco", "ameaça", "ameaca"]),
        "valor": _primeira_ocorrencia_apos(t, ["objetivo", "meta", "resultado esperado", "ROI"]),
    }
    return {"contexto": contexto, "mapa": mapa, "casos_uso": []}


def _achar_empresa(t: str) -> str:
    match = re.search(r"\b([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+)?)\s+(?:Ltda|S\.?A\.?|EIRELI|LTDA|Corp\.?|Inc\.?)\b", t)
    if match:
        return match.group(0)
    match = re.search(r"empresa\s+([A-Z][A-Za-z\s]{2,40})", t)
    if match:
        return match.group(1).strip()
    return ""


def _tem_pmo(t: str) -> bool:
    return bool(re.search(r"\b(PMO|escritório de projetos|escritorio de projetos)\b", t, re.IGNORECASE))


def _resumo_primeiras_frases(t: str, n: int = 3) -> str:
    frases = re.split(r"(?<=[.!?])\s+", t.strip())
    return " ".join(frases[:n])[:600]


def _primeira_ocorrencia_apos(t: str, keywords: list[str]) -> str:
    """Retorna a frase que contém uma das keywords, se houver."""
    frases = re.split(r"(?<=[.!?])\s+", t)
    for f in frases:
        f_lower = f.lower()
        if any(k in f_lower for k in keywords):
            return f.strip()[:400]
    return ""


def sugerir(texto: str, api_key: str | None = None) -> dict:
    """Dispatch: LLM (Groq) se chave disponível, senão heurística."""
    if api_key or api_key_disponivel():
        try:
            return sugerir_via_llm(texto, api_key=api_key)
        except ErroExtracao:
            return heuristica_fallback(texto)
    return heuristica_fallback(texto)


def sugestoes_para_casos_de_uso(sugestoes: dict) -> list[CasoDeUso]:
    """Converte a lista `casos_uso` das sugestões em objetos CasoDeUso prontos
    para injetar no session_state da aba 9."""
    casos = []
    for i, c in enumerate(sugestoes.get("casos_uso") or [], start=1):
        casos.append(
            CasoDeUso(
                id=c.get("id") or f"upload-{i}",
                nome=c.get("nome") or f"Caso importado {i}",
                contexto=c.get("contexto", ""),
                dor=c.get("dor", ""),
                dados=c.get("dados", ""),
                decisao=c.get("decisao", ""),
                dono=(c.get("dono") or None) or None,
                metrica_valor=c.get("metrica_valor", ""),
                notas={cid: 3 for cid in CRITERIO_IDS},
            )
        )
    return casos
